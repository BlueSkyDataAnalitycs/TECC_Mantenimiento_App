#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Informe Integral del Taller TECC v2 (.docx) — Prompt Maestro v2.0.
Genera el informe con branding TECC a partir del dataset ya computado por el app
(mismas tarifas/valorización que la plataforma, Regla de Oro #6). No inventa datos:
lo que no existe se marca N.D. y se reporta en el Módulo 6.
Uso: python3 informe_taller.py [ruta_data.json] [salida.docx]
"""
import json, sys, datetime, collections, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA = sys.argv[1] if len(sys.argv) > 1 else "/tmp/informe_data.json"
OUT = sys.argv[2] if len(sys.argv) > 2 else None

# ---- carga (el dump puede venir doble-codificado) ----
raw = open(DATA).read().strip()
d = json.loads(raw)
if isinstance(d, str):
    d = json.loads(d)
RECS = d["records"]
TABLA = d["tabla"]
COSTO_MES = 7_500_000                      # costo fijo mensual del taller (prompt)
FLOTA_CC = {"WG": 30, "TECC": 23, "CARLOS GASCA": 5, "NIDIA GASCA": 5, "404": 16, "(sin CC)": 3}
BRAYAN_MES = d.get("brayanMensual", 2_907_595)

# ---- marca ----
NAVY = RGBColor(0x1F, 0x38, 0x64)
GOLD = RGBColor(0xC9, 0xA2, 0x27)
VERDE = RGBColor(0x00, 0x66, 0x00)
ROJO = RGBColor(0xCC, 0x00, 0x00)
GRIS = RGBColor(0x66, 0x66, 0x66)
NAVY_HEX, GOLD_HEX, ZEBRA_HEX = "1F3864", "C9A227", "EEF1F6"

def cop(v):
    try: return "$ " + format(int(round(v)), ",d").replace(",", ".")
    except: return "N.D."

def pct(v): return ("%.1f%%" % (v * 100)).replace(".", ",")

# ---------- período ----------
fechas = sorted(r["fecha"] for r in RECS if r["fecha"])
f0 = datetime.date.fromisoformat(fechas[0]); f1 = datetime.date.fromisoformat(fechas[-1])
DIAS = (f1 - f0).days + 1
MESES = ["", "enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
def flarga(f): return "%d de %s de %d" % (f.day, MESES[f.month], f.year)
PERIODO = "%s – %s" % (flarga(f0), flarga(f1))
PERIODO_CORTO = "%d %s – %d %s %d" % (f0.day, MESES[f0.month][:3], f1.day, MESES[f1.month][:3], f1.year)
TOKEN = "%s-%s_%d" % (MESES[f0.month][:3].upper(), MESES[f1.month][:3].upper(), f1.year)
if OUT is None:
    OUT = "/Users/juliancho/Downloads/INFORME_TALLER_TECC_%s_v2.docx" % TOKEN
COSTO_PERIODO = COSTO_MES * DIAS / 30.0
HOY = flarga(datetime.date.fromisoformat(d.get("generated", "2026-07-04")[:10])) if d.get("generated") else flarga(f1)

# ================= DOCUMENTO =================
doc = Document()
st = doc.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(10.5)
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Inches(0.9); sec.left_margin = sec.right_margin = Inches(0.9)
CW = int(doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin)  # content width (EMU→ usamos como base para anchos por proporción)

def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)

def _set_cell(cell, text, bold=False, color=None, align="left", size=9.5, white=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.text = ""
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    run = p.add_run(str(text)); run.font.size = Pt(size); run.font.bold = bold; run.font.name = "Arial"
    if white: run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color is not None: run.font.color.rgb = color
    for m in ("top", "bottom"):
        pass

def table(headers, rows, widths=None, aligns=None, cell_colors=None):
    """headers: list; rows: list de list(str). aligns por columna. cell_colors: dict (i,j)->RGBColor."""
    n = len(headers)
    t = doc.add_table(rows=1, cols=n); t.style = "Table Grid"; t.autofit = False
    aligns = aligns or (["left"] + ["right"] * (n - 1))
    if widths:
        tot = sum(widths)
        for j, w in enumerate(widths):
            for cell in t.columns[j].cells: cell.width = int(CW * w / tot)
    for j, hdr in enumerate(headers):
        c = t.rows[0].cells[j]; _shade(c, NAVY_HEX); _set_cell(c, hdr, bold=True, align=aligns[j], white=True, size=9)
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            col = (cell_colors or {}).get((i, j))
            if i % 2 == 1 and col is None: _shade(cells[j], ZEBRA_HEX)
            _set_cell(cells[j], val, align=aligns[j], color=col, bold=(col is not None))
    doc.add_paragraph()
    return t

def h1(txt, num=None):
    doc.add_paragraph()
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(("MÓDULO %s — " % num if num else "") + txt); r.font.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY; r.font.name = "Arial"
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "18"); bot.set(qn("w:color"), GOLD_HEX); bot.set(qn("w:space"), "2")
    pbdr.append(bot); pPr.append(pbdr)

def h2(txt):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(txt); r.font.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = NAVY
    return p

def para(txt, size=10.5, italic=False, color=None, space=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space)
    r = p.add_run(txt); r.font.size = Pt(size); r.italic = italic
    if color is not None: r.font.color.rgb = color
    return p

def bullet(txt, size=10):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(txt); r.font.size = Pt(size); return p

def kpi_pill(label, value, color=NAVY):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
    r = p.add_run("● "); r.font.color.rgb = color; r.font.size = Pt(10)
    r2 = p.add_run(label + ":  "); r2.font.size = Pt(10)
    r3 = p.add_run(value); r3.font.bold = True; r3.font.size = Pt(10); r3.font.color.rgb = color

def barra(frac, ancho=22):
    llenos = max(0, min(ancho, int(round(frac * ancho))))
    return "█" * llenos + "░" * (ancho - llenos)

# footer con paginación + marca
def _footer():
    ft = doc.sections[0].footer; p = ft.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Transportes Especiales Ciudad de Cali S.A.S. — TECC   ·   Informe del taller %s   ·   Pág. " % PERIODO_CORTO)
    run.font.size = Pt(7.5); run.font.color.rgb = GRIS
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); r2 = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "15"); rpr.append(sz); r2.append(rpr); t = OxmlElement("w:t"); t.text = "1"; r2.append(t); fld.append(r2); p._p.append(fld)
_footer()

# ================= PORTADA =================
for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("TECC"); r.font.size = Pt(40); r.font.bold = True; r.font.color.rgb = NAVY
r2 = p.add_run("  ·  TALLER INTERNO"); r2.font.size = Pt(18); r2.font.color.rgb = GOLD; r2.font.bold = True
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Informe Integral del Taller"); r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = NAVY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Rentabilidad · Repuestos · Estado de Flota · Conductores · Operación"); r.font.size = Pt(12); r.font.color.rgb = GRIS
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Período analizado\n"); r.font.size = Pt(11); r.font.color.rgb = GOLD; r.font.bold = True
r = p.add_run(PERIODO); r.font.size = Pt(15); r.font.bold = True
for _ in range(6): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Transportes Especiales Ciudad de Cali S.A.S. — Grupo William Gasca\nTaller interno de mantenimiento (Taller Julián)\nGenerado el %s" % HOY)
r.font.size = Pt(10); r.font.color.rgb = GRIS
doc.add_page_break()

# ================= MÓDULO 1 — RENTABILIDAD =================
h1("RENTABILIDAD DEL TALLER", 1)
con_valor = [r for r in RECS if r["valor"] > 0]
valor_gen = sum(r["valor"] for r in con_valor)
sin_valor = len(RECS) - len(con_valor)
rent = valor_gen - COSTO_PERIODO
roi = rent / COSTO_PERIODO if COSTO_PERIODO else 0
cobertura = valor_gen / COSTO_PERIODO if COSTO_PERIODO else 0
proy_mes = valor_gen * 30.0 / DIAS

h2("1.1  Resumen ejecutivo del período")
kpi_pill("Valor generado (valorización de tercerización)", cop(valor_gen))
kpi_pill("Costo fijo del taller prorrateado (%d días)" % DIAS, cop(COSTO_PERIODO))
kpi_pill("Rentabilidad neta del período", cop(rent), VERDE if rent >= 0 else ROJO)
kpi_pill("ROI del período", pct(roi), VERDE if roi >= 0.30 else ROJO)
kpi_pill("Cobertura de costos", ("%.2fx" % cobertura).replace(".", ","), VERDE if cobertura >= 1.3 else ROJO)
kpi_pill("Proyección mensual (valor × 30/días)", cop(proy_mes))
para("Fórmula del costo: $7.500.000 mensuales × (%d días ÷ 30) = %s. Valorización según la tabla de tarifas de tercerización TECC (Módulo 1.2)." % (DIAS, cop(COSTO_PERIODO)), size=8.5, italic=True, color=GRIS)
if rent < 0:
    para("⚠️ Lectura de contexto: el resultado negativo compara %s meses de costo fijo contra solo las %d intervenciones registradas. Con un promedio de %.1f intervenciones/mes es muy probable que el taller esté SUB-REGISTRANDO trabajo (no que sea deficitario). La rentabilidad real solo será medible cuando se registre el 100%% de las intervenciones — ver Módulo 6." % (("%.1f" % (DIAS / 30)).replace(".", ","), len(RECS), len(RECS) / (DIAS / 30)), size=9, color=ROJO)

h2("1.2  Valorización de intervenciones vs. tercerización")
cat = collections.defaultdict(lambda: [0, 0])
for r in con_valor:
    c = r["categoria"] or "(sin categoría)"; cat[c][0] += 1; cat[c][1] += r["valor"]
filas = sorted(cat.items(), key=lambda kv: -kv[1][1])
rows = [[c, str(n), cop(TABLA.get(c, tot // n if n else 0)), cop(tot)] for c, (n, tot) in filas]
rows.append(["TOTAL", str(len(con_valor)), "", cop(valor_gen)])
table(["Servicio / categoría", "N°", "Tarifa unit.", "Valor tercerización"], rows,
      widths=[52, 10, 18, 20], aligns=["left", "center", "right", "right"],
      cell_colors={(len(rows) - 1, 3): NAVY})
para("%d intervenciones no se pudieron valorar con la tabla (descripción insuficiente); se listan en el Módulo 6." % sin_valor, size=8.5, italic=True, color=GRIS)

h2("1.3  Mix correctivo / preventivo")
tipos = collections.Counter((r["tipo"] or "N.D.").title() for r in RECS)
trows = [[k, str(v), pct(v / len(RECS))] for k, v in tipos.most_common()]
table(["Tipo de mantenimiento", "N°", "% del total"], trows, widths=[60, 20, 20], aligns=["left", "center", "right"])
para("⚠️ El tipo de mantenimiento solo viene diligenciado en %d de %d intervenciones (el resto = N.D.). Meta: elevar el preventivo; hoy el dato es insuficiente para medir el mix. Ver Módulo 6." % (sum(1 for r in RECS if r["tipo"]), len(RECS)), size=9, color=ROJO)

h2("1.4  Alertas de reincidencia (2+ visitas en el período)")
veh_ct = collections.Counter(r["placa"] for r in RECS if r["placa"])
reinc = [(v, n) for v, n in veh_ct.items() if n >= 2]
reinc.sort(key=lambda x: -x[1])
para("Vehículos con 2 o más ingresos al taller en el período: %d." % len(reinc))
if reinc:
    table(["N° interno", "Ingresos"], [[v, str(n)] for v, n in reinc[:15]], widths=[60, 40], aligns=["left", "center"])

h2("1.5  Prorrateo del costo fijo por centro de costo")
tot_flota = sum(FLOTA_CC.values())
gen_cc = collections.defaultdict(float)
for r in con_valor:
    gen_cc[r["centro"]] += r["valor"]
rows = []; ccc = {}
orden = sorted(FLOTA_CC, key=lambda c: -FLOTA_CC[c])
for i, c in enumerate(orden):
    nv = FLOTA_CC[c]; share = nv / tot_flota; asign = COSTO_PERIODO * share
    gen = gen_cc.get(c, 0.0); diff = gen - asign
    rows.append([c, str(nv), pct(share), cop(asign), cop(gen), ("+ " if diff >= 0 else "− ") + cop(abs(diff))])
    ccc[(i, 5)] = VERDE if diff >= 0 else ROJO
table(["Centro de costo (propietario)", "Veh.", "% flota", "Costo asignado", "Valor generado", "Diferencia"],
      rows, widths=[30, 8, 12, 18, 16, 16], aligns=["left", "center", "center", "right", "right", "right"], cell_colors=ccc)
para("El costo fijo se reparte por número de vehículos de cada propietario (%d en total). Verde = el taller propio le ahorra a ese CC; rojo = déficit. Centros «404» y «(sin CC)» = vehículos sin propietario claro en la base." % tot_flota, size=8.5, italic=True, color=GRIS)

doc.add_page_break()

# ================= MÓDULO 2 — REPUESTOS =================
h1("REPUESTOS", 2)
para("Estado del dato: NO DISPONIBLE.", color=ROJO)
para("Ni el formulario «Intervenciones Vehiculares TECC» actualmente desplegado ni el registro histórico del jefe de taller capturan repuestos ni su condición. Por lo tanto, los KPIs 4, 5 y 6 (cobertura de repuestos, tasa de usados/reencauchados y alerta de seguridad) quedan en N.D. este período.")
h2("Qué falta para activar este módulo")
bullet("Agregar al formulario la pregunta «¿Se instalaron repuestos y en qué condición?» con opciones: No se usaron / Todos nuevos / Incluye usados o reencauchados / Mixto.")
bullet("Agregar «Detalle de repuestos (referencia, marca, cantidad)» como texto libre para el ranking del Módulo 2.6.")
para("En cuanto esas dos columnas empiecen a llegar, este módulo se llena solo (cobertura, condición por centro de costo, alerta 🔴 de repuesto usado en frenos-suspensión-dirección, top consumidores y top repuestos).", size=9, italic=True, color=GRIS)

# ================= MÓDULO 3 — ESTADO DE FLOTA =================
h1("ESTADO DE LA FLOTA", 3)
est = [r for r in RECS if r["estado"] is not None]
para("Intervenciones con estado de ingreso (escala 1–10) diligenciado: %d de %d." % (len(est), len(RECS)), color=(ROJO if len(est) < 20 else NAVY))
if est:
    iei = sum(r["estado"] for r in est) / len(est)
    h2("3.1–3.2  Índice de Estado de Ingreso (IEI) y semáforo")
    kpi_pill("IEI del período (promedio 1–10)", ("%.1f / 10" % iei).replace(".", ","), VERDE if iei >= 7 else (GOLD if iei >= 5 else ROJO))
    bins = [("🔴 Crítico (1–4)", lambda e: e <= 4), ("🟡 Alerta (5–6)", lambda e: 5 <= e <= 6),
            ("🟢 Aceptable (7–8)", lambda e: 7 <= e <= 8), ("⭐ Óptimo (9–10)", lambda e: e >= 9)]
    rows = []
    for lab, fn in bins:
        n = sum(1 for r in est if fn(r["estado"])); fr = n / len(est)
        rows.append([lab, str(n), pct(fr), barra(fr)])
    table(["Franja", "N°", "%", "Distribución"], rows, widths=[26, 8, 12, 54], aligns=["left", "center", "right", "left"])
    crit = sum(1 for r in est if r["estado"] <= 4)
    para("Vehículos-intervención en estado crítico (≤4): %d (%s)." % (crit, pct(crit / len(est))), size=9)
    h2("3.4  Estado promedio por sistema intervenido (n≥2)")
    sysd = collections.defaultdict(list)
    for r in est: sysd[r["sistema"] or "(sin sistema)"].append(r["estado"])
    rows = [[s, str(len(v)), ("%.1f" % (sum(v) / len(v))).replace(".", ",")] for s, v in sorted(sysd.items(), key=lambda kv: sum(kv[1]) / len(kv[1])) if len(v) >= 2]
    if rows: table(["Sistema", "N°", "Estado prom."], rows, widths=[64, 12, 24], aligns=["left", "center", "center"])
    else: para("Aún no hay sistemas con n≥2 intervenciones con estado registrado.", size=9, italic=True, color=GRIS)
para("Nota: el histórico migrado no trae estado de ingreso; por eso el IEI se calcula solo sobre las intervenciones del formulario nuevo. La muestra crecerá a medida que el taller diligencie la escala 1–10.", size=8.5, italic=True, color=GRIS)

# ================= MÓDULO 4 — CONDUCTORES =================
h1("CONDUCTORES", 4)
para("Estado del dato: NO DISPONIBLE (Índice de Conductor).", color=ROJO)
para("El formulario actual no registra el conductor que entrega el vehículo, y el histórico solo trae la cédula (sin nombre). Por lo tanto el Índice de Conductor (ICE, KPI 9) y el Delta conductor–vehículo (KPI 10) no se pueden calcular este período.")
h2("Qué falta para activar este módulo")
bullet("Agregar al formulario «Conductor que entrega el vehículo (primer nombre y primer apellido)».")
bullet("Con eso se calcula el ICE por conductor (estado promedio de entrega, solo n≥3), el ranking peor→mejor y el delta contra el estado histórico de sus vehículos.")
h2("4.4  Nota metodológica (obligatoria)")
para("Correlación ≠ causalidad. El Índice de Conductor es un insumo de gestión y formación, NO una prueba disciplinaria. El estado con que un vehículo ingresa al taller depende también de su edad, ruta y kilometraje, no solo del conductor.", size=9.5, italic=True)

# ================= MÓDULO 5 — OPERACIÓN =================
h1("OPERACIÓN DEL TALLER", 5)
dur = [r for r in RECS if r["durMin"] is not None]
atip = [r for r in dur if r["durMin"] > 720]
normal = [r for r in dur if 0 < r["durMin"] <= 720]
para("Intervenciones con hora de inicio y fin válidas: %d (%d atípicas >12h excluidas de promedios)." % (len(dur), len(atip)), color=(ROJO if len(dur) < 15 else NAVY))
if normal:
    h2("5.1  Duración promedio de intervención")
    prom = sum(r["durMin"] for r in normal) / len(normal)
    kpi_pill("Duración media (sin atípicos)", ("%.0f min  (≈ %.1f h)" % (prom, prom / 60)).replace(".", ","))
    h2("5.2  Productividad por técnico")
    tecd = collections.defaultdict(lambda: [0, 0])
    for r in dur:
        t = (r["tecnico"] or "N.D.").strip().title()
        tecd[t][0] += 1
        if r["durMin"] and r["durMin"] <= 720: tecd[t][1] += r["durMin"]
    rows = [[t, str(n), ("%.1f h" % (mn / 60)).replace(".", ",")] for t, (n, mn) in sorted(tecd.items(), key=lambda kv: -kv[1][0])]
    table(["Técnico / mecánico", "Intervenciones", "Horas registradas"], rows, widths=[54, 22, 24], aligns=["left", "center", "right"])
    h2("5.3  Costo por hora efectiva del taller")
    horas_tot = sum(r["durMin"] for r in normal) / 60.0
    cxh = COSTO_PERIODO / horas_tot if horas_tot else 0
    kpi_pill("Costo por hora efectiva (costo prorrateado ÷ horas registradas)", cop(cxh))
    para("⚠️ Solo %d de %d intervenciones traen horas; este costo/hora es indicativo hasta que el formulario capture inicio y fin en todas." % (len(dur), len(RECS)), size=9, color=ROJO)
# 5.4 reincidencia técnica
h2("5.4  Reincidencia técnica (mismo vehículo + sistema ≤ 30 días)")
byvs = collections.defaultdict(list)
for r in RECS:
    if r["placa"] and r["sistema"] and r["fecha"]:
        byvs[(r["placa"], r["sistema"])].append(datetime.date.fromisoformat(r["fecha"]))
reint = 0
for k, ds in byvs.items():
    ds.sort()
    for a, b in zip(ds, ds[1:]):
        if (b - a).days <= 30: reint += 1
para("Casos de reincidencia técnica (posible retrabajo o repuesto de mala calidad): %d." % reint)
h2("5.5  Cumplimiento de evidencia")
para("El formulario captura VIDEO previo y posterior (no foto). La cobertura de evidencia audiovisual no se consolidó en esta versión de la extracción → N.D. (Módulo 6).", size=9, italic=True, color=GRIS)

doc.add_page_break()

# ================= KPIs v2 =================
h1("TABLERO DE KPIs v2")
def sem(ok): return ("🟢", VERDE) if ok else ("🔴", ROJO)
kpis = []
s, c = sem(rent > 0); kpis.append(["1", "Rentabilidad neta", cop(rent), s])
s, c = sem(roi >= 0.30); kpis.append(["2", "ROI del período (meta ≥30%)", pct(roi), s])
s, c = sem(cobertura >= 1.3); kpis.append(["3", "Cobertura de costos (meta ≥1,3x)", ("%.2fx" % cobertura).replace(".", ","), s])
kpis.append(["4", "% intervenciones con repuestos", "N.D.", "⚪"])
kpis.append(["5", "Tasa de repuestos usados (meta <15%)", "N.D.", "⚪"])
kpis.append(["6", "Alerta seguridad repuestos (meta 0)", "N.D.", "⚪"])
if est:
    s, c = sem(iei >= 7); kpis.append(["7", "Índice de Estado de Ingreso (IEI)", ("%.1f/10 (n=%d)" % (iei, len(est))).replace(".", ","), s if iei >= 7 else "🟡"])
else:
    kpis.append(["7", "Índice de Estado de Ingreso (IEI)", "N.D.", "⚪"])
kpis.append(["8", "% flota crítica (meta <10%)", (pct(crit / len(est)) + " (n=%d)" % len(est)) if est else "N.D.", "⚪"])
kpis.append(["9", "Índice de Conductor (ICE)", "N.D.", "⚪"])
kpis.append(["10", "Delta conductor–vehículo", "N.D.", "⚪"])
kpis.append(["11", "Duración media de intervención", (("%.0f min" % prom).replace(".", ",") + " (n=%d)" % len(normal)) if normal else "N.D.", "⚪"])
kpis.append(["12", "Costo por hora efectiva", (cop(cxh) + "/h") if normal and horas_tot else "N.D.", "⚪"])
mixc = sum(1 for r in RECS if r["tipo"])
kpis.append(["13", "Mix correctivo", ("%d/%d con dato" % (mixc, len(RECS))), "⚪"])
kpis.append(["14", "Reincidencia 30 días (meta <8%)", "%d casos" % reint, "⚪"])
kpis.append(["15", "Cumplimiento fotográfico (meta ≥90%)", "N.D. (form usa video)", "⚪"])
table(["#", "KPI", "Valor del período", "Sem."], [[k[0], k[1], k[2], k[3]] for k in kpis],
      widths=[6, 54, 30, 10], aligns=["center", "left", "right", "center"])
para("⚪ = dato no disponible este período (ver Módulos 2, 4, 5 y 6). Los KPIs de estado usan la muestra del formulario nuevo (n indicado).", size=8.5, italic=True, color=GRIS)

# ================= MÓDULO 6 — CALIDAD DE DATOS =================
h1("CALIDAD DE DATOS", 6)
h2("Cobertura de campos (238 intervenciones del período)")
cov = lambda k, test=lambda x: str(x).strip(): sum(1 for r in RECS if test(r.get(k)))
campos = [("Fecha", cov("fecha")), ("N° interno (vehículo)", cov("placa")),
          ("Descripción del trabajo", sum(1 for r in RECS if r["categoria"] or r["sistema"])),
          ("Sistema clasificado", sum(1 for r in RECS if r["sistema"])),
          ("Valorizable con tabla de tarifas", len(con_valor)),
          ("Estado de ingreso (1–10)", len(est)), ("Técnico encargado", sum(1 for r in RECS if r["tecnico"])),
          ("Tipo (prev./corr.)", sum(1 for r in RECS if r["tipo"])), ("Hora inicio+fin (duración)", len(dur)),
          ("Kilometraje", sum(1 for r in RECS if r["km"])), ("Repuestos y condición", 0),
          ("Conductor (nombre)", 0)]
rows = [[c, "%d / %d" % (n, len(RECS)), pct(n / len(RECS)), barra(n / len(RECS))] for c, n in campos]
table(["Campo", "Con dato", "%", "Cobertura"], rows, widths=[36, 16, 12, 36], aligns=["left", "center", "right", "left"])
h2("Hallazgos y N.D. del período")
bullet("El formulario desplegado NO captura repuestos, condición de repuestos, componentes ni conductor → Módulos 2 y 4 en N.D.")
bullet("Estado de ingreso (1–10) solo en %d de %d intervenciones (el histórico migrado no lo trae)." % (len(est), len(RECS)))
bullet("Técnico y tipo de mantenimiento solo desde el 2-jul (formulario) → %d intervenciones." % sum(1 for r in RECS if r["tecnico"]))
bullet("%d intervenciones sin valorar (descripción insuficiente para la tabla de tarifas)." % sin_valor)
if atip: bullet("%d intervención(es) con duración atípica (>12 h): se muestran pero se excluyen de los promedios (posible turno cruzado o error de digitación de hora)." % len(atip))
bullet("Se excluyó 1 registro de PRUEBA del administrador (19-jun) en la extracción de aseos; el taller no tenía registros de prueba.")
bullet("Fuente única este período: no hubo cruce A↔B por placa+fecha (el registro histórico y el formulario se consolidan en una sola base).")
h2("Recomendación de datos (para desbloquear el informe completo v2)")
para("Volver obligatorias en el formulario: (1) Conductor que entrega, (2) Repuestos + condición, (3) Estado 1–10, (4) Técnico y (5) Tipo. Con esos 5 campos, los Módulos 2–5 y los KPIs 4–15 se llenan automáticamente en la próxima corrida.", size=9.5)

doc.save(OUT)
print("OK →", OUT)
print("Período: %s (%d días) · valor %s · costo %s · rentabilidad %s · ROI %s" % (PERIODO_CORTO, DIAS, cop(valor_gen), cop(COSTO_PERIODO), cop(rent), pct(roi)))
print("estado n=%d · tecnico n=%d · dur n=%d · reincidencias(2+)=%d · reincidencia30=%d" % (len(est), sum(1 for r in RECS if r["tecnico"]), len(dur), len(reinc), reint))
